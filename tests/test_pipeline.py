from pathlib import Path

import docx
import pytest
from reportlab.pdfgen import canvas

from ats_xray.pipeline import AnalysisResult, analyze_bytes, analyze_path, extract_text


def _make_two_column_pdf(path) -> None:
    """A section header split across two columns: readable layout-aware,
    invisible under naive extraction — the same fixture shape used by the
    Day 1/3/4 smoke tests.
    """
    c = canvas.Canvas(str(path), pagesize=(500, 300))
    c.setFont("Helvetica", 12)
    c.drawString(30, 270, "Jane Doe")
    c.drawString(30, 250, "Experience")
    c.drawString(280, 270, "jane@example.com")
    c.drawString(280, 250, "+1 555 123 4567")
    c.save()


def test_extract_text_pdf(tmp_path):
    pdf_path = tmp_path / "resume.pdf"
    _make_two_column_pdf(pdf_path)

    naive_text, aware_text = extract_text(str(pdf_path))

    assert "Jane Doe" in naive_text
    assert "Jane Doe" in aware_text
    assert naive_text != aware_text


def test_extract_text_docx(tmp_path):
    path = tmp_path / "resume.docx"
    document = docx.Document()
    document.add_paragraph("Jane Doe")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Skills"
    table.rows[0].cells[1].text = "Python"
    document.save(str(path))

    naive_text, aware_text = extract_text(str(path))

    assert "Skills" not in naive_text
    assert "Skills" in aware_text


def test_extract_text_unsupported_extension_raises(tmp_path):
    path = tmp_path / "resume.txt"
    path.write_text("hello")

    with pytest.raises(ValueError):
        extract_text(str(path))


def test_analyze_path_includes_findings(tmp_path):
    pdf_path = tmp_path / "resume.pdf"
    _make_two_column_pdf(pdf_path)

    result = analyze_path(str(pdf_path))

    assert isinstance(result, AnalysisResult)
    assert any(f.rule.id == "section_missing_under_naive_parsing" for f in result.findings)


def test_analyze_bytes_end_to_end(tmp_path):
    pdf_path = tmp_path / "resume.pdf"
    _make_two_column_pdf(pdf_path)
    file_bytes = pdf_path.read_bytes()

    result = analyze_bytes(file_bytes, "resume.pdf")

    assert "Jane Doe" in result.aware_text
    assert any(f.rule.id == "section_missing_under_naive_parsing" for f in result.findings)


def test_analyze_bytes_unsupported_extension_raises():
    with pytest.raises(ValueError):
        analyze_bytes(b"hello", "resume.txt")


def test_analyze_bytes_cleans_up_temp_file_even_on_failure(monkeypatch):
    """The temp file must be deleted even when analysis blows up midway —
    e.g. a corrupted upload — not just on the happy path.
    """
    import ats_xray.pipeline as pipeline_module

    captured_paths = []

    def failing_analyze_path(path):
        captured_paths.append(path)
        raise RuntimeError("boom")

    monkeypatch.setattr(pipeline_module, "analyze_path", failing_analyze_path)

    with pytest.raises(RuntimeError):
        analyze_bytes(b"not a real pdf", "resume.pdf")

    assert captured_paths
    assert not Path(captured_paths[0]).exists()
