"""Five CVs, worst to best, each written against one of the vacancies.

The rest of the fixtures in this suite each isolate one fault, which is
what makes them good unit fixtures and poor evidence: a real CV is never
one problem, and the interesting question is whether the tool ranks two
imperfect documents the way a person would.

So these are graded rather than labelled. Each carries a plausible mix of
what real CVs get wrong, and the test asserts the *order*: parse readiness
rises from the first to the fifth. An individual number is a judgement
call; the ordering is not, and it is what breaks when a detector
regresses.

The match score is deliberately not ordered with them. Each CV is measured
against a different advert, so comparing those numbers across the corpus
compares two different questions — and the first CV out-scoring the second
is a real result worth keeping: a document can be barely readable and hold
exactly the right words, while a tidy one misses the licence the advert
insisted on.

Each function writes one file and returns nothing. Not a test module (no
test_ prefix), so pytest does not collect it.
"""

import docx
from reportlab.pdfgen import canvas

A4 = (595, 842)


def _page(path, lines, size=A4, font="Helvetica", start=790, leading=17, left=57):
    """One plain page of text, the shape most CVs actually are."""
    c = canvas.Canvas(str(path), pagesize=size)
    c.setFont(font, 11)
    y = start
    for line in lines:
        c.drawString(left, y, line)
        y -= leading
    c.save()


# --------------------------------------------------------------------------
# 1 — everything wrong at once
# --------------------------------------------------------------------------


def worst_care_docx(path) -> None:
    """A care CV built entirely as a Word table, reachable only by a link.

    The shape a template produces when someone lays out a CV visually: two
    columns held together by a table, the contact line a clickable profile,
    the section labels written as prose. Every one of those is invisible or
    scrambled to a parser, and none of them looks wrong on screen.

    The ligature in "Pflegefachkraft" is the kind a design tool inserts on
    export without asking.
    """
    document = docx.Document()
    document.add_paragraph("Marina Kowalski")
    document.add_paragraph("LinkedIn: linkedin.com/in/marinakowalski")

    table = document.add_table(rows=4, cols=2)
    table.cell(0, 0).text = "Mein Weg"
    table.cell(0, 1).text = "2019 - 2024\nSeniorenzentrum Weser, Bremen\nPﬂegefachkraft im Wohnbereich"
    table.cell(1, 0).text = "Davor"
    table.cell(1, 1).text = "2015 - 2019\nAmbulanter Dienst Nord\nGrundpflege und Dokumentation"
    table.cell(2, 0).text = "Was ich kann"
    table.cell(2, 1).text = "Grundpflege, Medikamentengabe, Geduld"
    table.cell(3, 0).text = "Womit ich angefangen habe"
    table.cell(3, 1).text = "2012 - 2015\nAusbildung Altenpflege, Bremen"

    document.save(str(path))


# --------------------------------------------------------------------------
# 2 — readable to a person, half of it lost to a parser
# --------------------------------------------------------------------------


def poor_logistics_pdf(path) -> None:
    """A warehouse CV in two columns, with the contact line in a footer.

    Nothing here is careless: the layout is tidy and the content is real.
    It fails because the second column is read as a continuation of the
    first, and because the only phone number repeats at the foot of both
    pages, where parsers strip boilerplate.

    The forklift licence the advert insists on is missing, and the
    warehouse software sits in a job that ended over a decade ago.
    """
    c = canvas.Canvas(str(path), pagesize=A4)
    for page in range(2):
        c.setFont("Helvetica", 11)
        if page == 0:
            c.drawString(57, 790, "Tomasz Nowak")
            c.drawString(57, 760, "Berufserfahrung")
            c.drawString(57, 740, "2018 - heute")
            c.drawString(57, 723, "Logistikzentrum Rhein, Monheim")
            c.drawString(57, 706, "Kommissionierung und Warenverraeumung")
            c.drawString(57, 676, "2006 - 2012")
            c.drawString(57, 659, "Spedition Alt, Duisburg")
            c.drawString(57, 642, "Lagerverwaltung mit SAP, Inventur")
            # The second column: a person reads it as a sidebar, a
            # layout-blind parser reads it as the same lines continued.
            c.drawString(330, 790, "Kenntnisse")
            c.drawString(330, 770, "Kommissionierung")
            c.drawString(330, 753, "Wareneingang")
            c.drawString(330, 736, "Zwei-Schicht-Betrieb")
        else:
            c.drawString(57, 790, "Ausbildung")
            c.drawString(57, 770, "2003 - 2006")
            c.drawString(57, 753, "Fachkraft fuer Lagerlogistik, Duisburg")
        # The only way to reach him, in the place a parser discards.
        c.setFont("Helvetica", 8)
        c.drawString(57, 40, "Tomasz Nowak | 0211 4455667 | Monheim am Rhein")
        c.showPage()
    c.save()


# --------------------------------------------------------------------------
# 3 — ordinary, with an ordinary gap
# --------------------------------------------------------------------------


def middling_retail_pdf(path) -> None:
    """A retail CV that parses cleanly and says too little.

    Single column, headings a parser knows, dates it can read. What is
    missing is on the content side: no email, and the till work the advert
    is mostly about is named once in passing rather than described.

    This is the common case, and the one where a parse-readiness score and
    a match score disagree: the file is fine and the application is thin.
    """
    _page(
        path,
        [
            "Sabine Reuter",
            "0176 2233445",
            "",
            "Berufserfahrung",
            "03/2021 - heute   Getraenkemarkt Sued, Hannover",
            "Verkauf und Beratung, Regalpflege",
            "",
            "09/2017 - 02/2021   Baeckerei Lindemann, Hannover",
            "Verkauf und Kassiertaetigkeit",
            "",
            "Ausbildung",
            "08/2014 - 07/2017   Realschulabschluss, Hannover",
            "",
            "Kenntnisse",
            "Kundenberatung, Zuverlaessigkeit, Deutsch als Muttersprache",
        ],
    )


# --------------------------------------------------------------------------
# 4 — a good CV with one thing left on the table
# --------------------------------------------------------------------------


def good_design_pdf(path) -> None:
    """A design CV that does everything right except name one skill.

    Contact in plain text on line two, conventional headings, months in the
    dates, the software the advert asks for named where it was used. The
    only gap is a preference the advert marked as nice-to-have, which is
    what the "what to fix first" list should pick up and weigh accordingly.
    """
    _page(
        path,
        [
            "Anna Muster",
            "anna.muster@example.com | +49 170 1234567 | Hamburg",
            "",
            "Berufserfahrung",
            "Maerz 2019 - heute   Studio Nord, Hamburg",
            "Mediengestalterin: Kataloge, Anzeigen und Verpackungen",
            "Reinzeichnung und Druckvorstufe, Uebergabe an die Druckerei",
            "InDesign, Photoshop und Illustrator im taeglichen Einsatz",
            "",
            "Januar 2016 - Februar 2019   Agentur Elbwerk, Hamburg",
            "Layout und Typografie fuer Printkampagnen",
            "Bildbearbeitung und Freisteller, Figma fuer Abstimmungen",
            "",
            "Ausbildung",
            "September 2012 - August 2015   Ausbildung Mediengestalterin",
            "Digital und Print, Hamburg",
            "",
            "Kenntnisse",
            "InDesign, Photoshop, Illustrator, Figma, Typografie, Layout",
            "",
            "Sprachen",
            "Deutsch C2, Englisch B2",
        ],
    )


# --------------------------------------------------------------------------
# 5 — nothing left for the tool to say
# --------------------------------------------------------------------------


def best_software_pdf(path) -> None:
    """A software CV with no parsing risk and no unmet requirement.

    The control at the top end. If this one ever picks up a finding, the
    finding is wrong: there is no table, no column, no image, no header,
    both contact fields are plain text on line two, every heading is one
    the parser knows, and every date carries a month.
    """
    _page(
        path,
        [
            "Daniel Weiss",
            "daniel.weiss@example.com | +49 151 9876543 | Koeln",
            "",
            "Berufserfahrung",
            "April 2018 - heute   sync.blue GmbH, remote",
            "Senior Software Engineer: Python und Django im Backend",
            "Betrieb auf Kubernetes, Docker und Linux",
            "PostgreSQL, REST-APIs und CI/CD, Verantwortung fuer Releases",
            "Terraform fuer die Infrastruktur, Code Reviews im Team",
            "",
            "Oktober 2014 - Maerz 2018   Meditec GmbH, Koeln",
            "Softwareentwickler: Python, Django, PostgreSQL",
            "Vue.js im Frontend",
            "",
            "Ausbildung",
            "Oktober 2010 - September 2014   Studium Informatik",
            "Bachelor of Science, Universitaet zu Koeln",
            "",
            "Kenntnisse",
            "Python, Django, PostgreSQL, Docker, Kubernetes, Linux,",
            "CI/CD, Terraform, Vue.js, REST",
            "",
            "Sprachen",
            "Deutsch C2, Englisch C1",
        ],
    )


GRADED = (
    ("1_worst_care", "pflegefachkraft", worst_care_docx, ".docx"),
    ("2_poor_logistics", "lagerlogistik", poor_logistics_pdf, ".pdf"),
    ("3_middling_retail", "verkauf", middling_retail_pdf, ".pdf"),
    ("4_good_design", "mediengestalter", good_design_pdf, ".pdf"),
    ("5_best_software", "python_engineer", best_software_pdf, ".pdf"),
)
"""The corpus in order, each CV with the advert it was written against."""


def write_all(directory) -> list:
    """Write every CV and copy every advert into ``directory``.

    So the corpus can be looked at rather than only asserted about: run
    ``python tests/corpus_generators.py <folder>`` and upload the files to
    the app. A fixture nobody can open is a fixture nobody checks.
    """
    from pathlib import Path
    from shutil import copyfile

    directory = Path(directory)
    directory.mkdir(parents=True, exist_ok=True)
    vacancies = Path(__file__).parent / "corpus" / "vacancies"

    written = []
    for name, vacancy, build, suffix in GRADED:
        path = directory / f"{name}{suffix}"
        build(path)
        written.append(path)
        advert = directory / f"{name}__advert.txt"
        copyfile(vacancies / f"{vacancy}.txt", advert)
        written.append(advert)
    return written


if __name__ == "__main__":
    import sys

    target = sys.argv[1] if len(sys.argv) > 1 else "corpus_out"
    for written in write_all(target):
        print(written)
