import docx
from docx.oxml import parse_xml
from docx.oxml.ns import nsmap

from ats_xray.docx_extract import extract_docx_full, extract_docx_naive
from ats_xray.docx_structure import extract_docx_headers_footers, find_docx_text_box_content, has_table_content


def test_extract_docx_headers_footers_reads_distinct_content(tmp_path):
    path = tmp_path / "resume.docx"
    document = docx.Document()
    document.sections[0].header.paragraphs[0].text = "John Smith | john@example.com"
    document.sections[0].footer.paragraphs[0].text = "Confidential"
    document.add_paragraph("Body content.")
    document.save(str(path))

    result = extract_docx_headers_footers(str(path))

    assert result["headers"] == ["John Smith | john@example.com"]
    assert result["footers"] == ["Confidential"]


def test_extract_docx_headers_footers_empty_when_absent(tmp_path):
    path = tmp_path / "resume.docx"
    docx.Document().save(str(path))

    assert extract_docx_headers_footers(str(path)) == {"headers": [], "footers": []}


def _document_with_text_box(path: str) -> None:
    document = docx.Document()
    document.add_paragraph("Normal paragraph, visible everywhere.")
    txbx_xml = (
        f'<w:txbxContent xmlns:w="{nsmap["w"]}">'
        "<w:p><w:r><w:t>Skills: Python, SQL, Docker</w:t></w:r></w:p>"
        "</w:txbxContent>"
    )
    document.element.body.append(parse_xml(txbx_xml))
    document.save(path)


def test_find_docx_text_box_content_finds_nested_text(tmp_path):
    path = tmp_path / "resume.docx"
    _document_with_text_box(str(path))

    assert find_docx_text_box_content(str(path)) == ["Skills: Python, SQL, Docker"]


def test_find_docx_text_box_content_empty_when_absent(tmp_path):
    path = tmp_path / "resume.docx"
    docx.Document().save(str(path))

    assert find_docx_text_box_content(str(path)) == []


def test_text_box_content_invisible_to_day_one_extractors(tmp_path):
    path = tmp_path / "resume.docx"
    _document_with_text_box(str(path))

    assert "Skills" not in extract_docx_naive(str(path))
    assert "Skills" not in extract_docx_full(str(path))


def test_has_table_content_true_when_cell_has_text(tmp_path):
    path = tmp_path / "resume.docx"
    document = docx.Document()
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Skills"
    table.rows[0].cells[1].text = "Python, SQL"
    document.save(str(path))

    assert has_table_content(str(path)) is True


def test_has_table_content_false_when_no_table(tmp_path):
    path = tmp_path / "resume.docx"
    docx.Document().save(str(path))

    assert has_table_content(str(path)) is False


def test_has_table_content_false_when_table_cells_are_empty(tmp_path):
    path = tmp_path / "resume.docx"
    document = docx.Document()
    document.add_table(rows=1, cols=2)
    document.save(str(path))

    assert has_table_content(str(path)) is False
