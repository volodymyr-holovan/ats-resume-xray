import subprocess
import sys

import docx

from ats_xray.cli import _format_field_comparison, _format_rule_report, _format_structure_report
from ats_xray.engine import Finding
from ats_xray.field_report import build_field_report
from ats_xray.rule import get_rule


def test_format_structure_report_pdf_with_findings():
    findings = {
        "non_embedded_fonts": ["Calibri"],
        "repeated_header_footer_lines": [{"zone": "footer", "text": "Page # of #", "pages": [1, 2]}],
        "textless_images": [{"page": 1, "bbox": (0, 0, 10, 10), "area_fraction": 0.2}],
    }

    report = _format_structure_report(findings)

    assert "Calibri" in report
    assert "footer" in report
    assert "20%" in report


def test_format_structure_report_pdf_no_findings():
    findings = {
        "non_embedded_fonts": [],
        "repeated_header_footer_lines": [],
        "textless_images": [],
    }

    report = _format_structure_report(findings)

    assert report.count("none found") == 3


def test_format_structure_report_docx():
    findings = {
        "headers_footers": {"headers": ["John Smith"], "footers": []},
        "text_box_content": ["Skills: Python"],
    }

    report = _format_structure_report(findings)

    assert "John Smith" in report
    assert "Skills: Python" in report
    assert "Footer content: none found" in report


def test_format_field_comparison_flags_field_at_risk_under_naive_parsing():
    aware_text = "Jane Doe\njane@example.com\n\nExperience\nSenior Engineer"
    naive_text = "Jane Doe Experience Senior Engineer jane@example.com"

    aware_report = build_field_report(aware_text)
    naive_report = build_field_report(naive_text)
    report = _format_field_comparison(aware_report, naive_report)

    assert "experience: layout-aware=found, naive=MISSING" in report
    assert "at risk" in report


def test_format_field_comparison_no_risk_when_both_agree():
    text = "Jane Doe\njane@example.com\n\nSkills\nPython"
    field_report = build_field_report(text)

    report = _format_field_comparison(field_report, field_report)

    assert "at risk" not in report
    assert "name: layout-aware=found, naive=found" in report


def test_format_rule_report_no_findings():
    assert _format_rule_report([]) == "No rules triggered."


def test_format_rule_report_includes_evidence_and_source():
    finding = Finding(rule=get_rule("pdf_non_embedded_font"), evidence_key="evidence_fonts", evidence_params={"fonts": "Calibri"})

    report = _format_rule_report([finding])

    assert "[MEDIUM] pdf_non_embedded_font" in report
    assert "Evidence: Non-embedded fonts: Calibri" in report
    assert "Source: research_sources.md#ats-fonts" in report


def test_format_rule_report_orders_high_severity_first():
    medium_finding = Finding(rule=get_rule("pdf_non_embedded_font"), evidence_key="evidence_fonts", evidence_params={"fonts": "Calibri"})
    high_finding = Finding(rule=get_rule("missing_contact_field"), evidence_key="evidence_no_contact")

    report = _format_rule_report([medium_finding, high_finding])

    assert report.index("missing_contact_field") < report.index("pdf_non_embedded_font")


def test_cli_handles_unicode_resume_content_without_crashing(tmp_path):
    """Regression test: cli.py used to inherit the OS console's default
    codepage for stdout, which crashed with UnicodeEncodeError on
    non-ASCII resume content (German umlauts, em-dashes) whenever that
    codepage couldn't represent it. main() now forces UTF-8 stdout/stderr
    regardless of locale — this runs the real CLI entry point as a
    subprocess to verify that end to end.
    """
    path = tmp_path / "resume.docx"
    document = docx.Document()
    document.add_paragraph("Jörg Müller")
    document.add_paragraph("jorg.mueller@beispiel.de | +49 151 23456789")
    document.add_paragraph("Straße: Königsallee 5, 40212 Düsseldorf")
    document.save(str(path))

    result = subprocess.run(
        [sys.executable, "-m", "ats_xray.cli", str(path)],
        capture_output=True,
        text=True,
        encoding="utf-8",
    )

    assert result.returncode == 0, result.stderr
    assert "Jörg Müller" in result.stdout


def test_cli_shows_friendly_message_for_missing_file(tmp_path):
    """Regression test: the CLI used to let FileNotFoundError propagate as
    a raw traceback. Should now exit with a clear one-line message.
    """
    missing_path = tmp_path / "does-not-exist.pdf"

    result = subprocess.run(
        [sys.executable, "-m", "ats_xray.cli", str(missing_path)],
        capture_output=True,
        text=True,
        encoding="utf-8",
    )

    assert result.returncode != 0
    assert "Traceback" not in result.stderr
    assert "not found" in result.stderr.lower()


def test_cli_shows_friendly_message_for_corrupted_file(tmp_path):
    """Regression test: a file with a .pdf name but invalid content used
    to let pdfminer's exception propagate as a raw traceback.
    """
    corrupt_path = tmp_path / "corrupt.pdf"
    corrupt_path.write_text("this is not a real pdf file at all")

    result = subprocess.run(
        [sys.executable, "-m", "ats_xray.cli", str(corrupt_path)],
        capture_output=True,
        text=True,
        encoding="utf-8",
    )

    assert result.returncode != 0
    assert "Traceback" not in result.stderr
    assert "couldn't read" in result.stderr.lower()
