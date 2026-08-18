"""Generates labeled resume fixtures, one per documented parsing-risk
pattern the rule engine knows about (plus one clean control). Each
function writes a single fixture file to the given path.

Used by test_golden_fixtures.py to build known-good and known-broken
resumes and assert the rule engine reacts to each pattern — no more, no
less — so a future change that silently breaks a detector shows up as a
failing test instead of a quiet regression.

Not a test module itself (no test_ prefix), so pytest won't try to collect
it directly.
"""

from pathlib import Path

import docx
from docx.oxml import parse_xml
from docx.oxml.ns import nsmap
from PIL import Image
from reportlab.pdfgen import canvas


def clean_single_column(path) -> None:
    """A well-formed, single-column resume with every expected section and
    contact field present. Nothing should trigger.
    """
    c = canvas.Canvas(str(path), pagesize=(400, 420))
    c.setFont("Helvetica", 12)
    c.drawString(30, 390, "Jane Doe")
    c.drawString(30, 370, "jane@example.com | +1 555 123 4567")
    c.drawString(30, 330, "Experience")
    c.drawString(30, 310, "Senior Engineer at Acme")
    c.drawString(30, 270, "Education")
    c.drawString(30, 250, "BSc Computer Science")
    c.drawString(30, 210, "Skills")
    c.drawString(30, 190, "Python, SQL")
    c.save()


def two_column_pdf(path) -> None:
    """A section header split across two columns: readable layout-aware,
    invisible under naive, order-blind extraction.
    """
    c = canvas.Canvas(str(path), pagesize=(500, 300))
    c.setFont("Helvetica", 12)
    c.drawString(30, 270, "Jane Doe")
    c.drawString(30, 250, "Experience")
    c.drawString(280, 270, "jane@example.com")
    c.drawString(280, 250, "+1 555 123 4567")
    c.save()


def missing_contact(path) -> None:
    """A resume with no email or phone number anywhere in the document."""
    c = canvas.Canvas(str(path), pagesize=(400, 300))
    c.setFont("Helvetica", 12)
    c.drawString(30, 270, "Jane Doe")
    c.drawString(30, 230, "Experience")
    c.drawString(30, 210, "Senior Engineer at Acme")
    c.save()


def pdf_textless_image(path) -> None:
    """A name banner exported as a picture instead of real text, with the
    rest of the resume as normal, readable text.
    """
    image_path = Path(path).with_suffix(".png")
    Image.new("RGB", (300, 80), color="white").save(image_path)

    c = canvas.Canvas(str(path), pagesize=(400, 300))
    c.drawImage(str(image_path), 30, 220, width=300, height=60)
    c.setFont("Helvetica", 10)
    c.drawString(30, 190, "jane@example.com | +1 555 123 4567")
    c.drawString(30, 150, "Experience")
    c.drawString(30, 130, "Senior Engineer at Acme")
    c.save()
    image_path.unlink()


def pdf_repeated_header_footer(path) -> None:
    """A footer with contact info repeating verbatim across two pages —
    single-column body text throughout, so this isolates the
    repeated-footer signal without also tripping the column-mangling one.
    """
    c = canvas.Canvas(str(path), pagesize=(400, 300))
    for page_num in (1, 2):
        c.setFont("Helvetica", 12)
        c.drawString(30, 270, "Jane Doe")
        c.drawString(30, 250, "Experience" if page_num == 1 else "Education")
        c.drawString(30, 150, f"Body content page {page_num}")
        c.drawString(30, 20, "jane@example.com | +1 555 123 4567")
        if page_num == 1:
            c.showPage()
    c.save()


def docx_with_table(path) -> None:
    """Resume content placed inside a DOCX table."""
    document = docx.Document()
    document.add_paragraph("Jane Doe")
    document.add_paragraph("jane@example.com, +1 555 123 4567")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Skills"
    table.rows[0].cells[1].text = "Python, SQL"
    document.save(str(path))


def docx_contact_in_header(path) -> None:
    """Contact info that lives only in the header: invisible to body-only
    extraction, and genuinely unreachable since it appears nowhere else.
    """
    document = docx.Document()
    document.sections[0].header.paragraphs[0].text = "jane@example.com | +1 555 123 4567"
    document.add_paragraph("Jane Doe")
    document.add_paragraph("Experience")
    document.add_paragraph("Senior Engineer at Acme")
    document.save(str(path))


def docx_text_box(path) -> None:
    """Resume content placed inside a Word text box."""
    document = docx.Document()
    document.add_paragraph("Jane Doe")
    document.add_paragraph("jane@example.com, +1 555 123 4567")
    document.add_paragraph("Experience")
    document.add_paragraph("Senior Engineer at Acme")
    txbx_xml = (
        f'<w:txbxContent xmlns:w="{nsmap["w"]}">'
        "<w:p><w:r><w:t>Skills: Python, SQL, Docker</w:t></w:r></w:p>"
        "</w:txbxContent>"
    )
    document.element.body.append(parse_xml(txbx_xml))
    document.save(str(path))
