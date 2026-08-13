import docx

from ats_xray.docx_extract import extract_docx_full, extract_docx_naive


def make_resume_with_table(path):
    document = docx.Document()
    document.add_paragraph("Jane Doe")
    document.add_paragraph("Summary paragraph visible to both extractors.")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Skills"
    table.rows[0].cells[1].text = "Python, SQL"
    document.save(path)


def test_naive_extraction_misses_table_content(tmp_path):
    path = tmp_path / "resume.docx"
    make_resume_with_table(str(path))

    naive_text = extract_docx_naive(str(path))

    assert "Jane Doe" in naive_text
    assert "Skills" not in naive_text
    assert "Python, SQL" not in naive_text


def test_full_extraction_includes_table_content(tmp_path):
    path = tmp_path / "resume.docx"
    make_resume_with_table(str(path))

    full_text = extract_docx_full(str(path))

    assert "Jane Doe" in full_text
    assert "Skills" in full_text
    assert "Python, SQL" in full_text
