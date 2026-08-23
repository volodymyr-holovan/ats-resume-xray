import docx
import pytest
from reportlab.pdfgen import canvas

from ats_xray.docx_render import convert_docx_to_pdf, find_soffice
from ats_xray.engine import attach_docx_regions, run_rules
from ats_xray.field_report import build_field_report
from ats_xray.pdf_locate import find_text_regions
from ats_xray.pipeline import analyze_path, extract_text
from ats_xray.structure import analyze_structure

needs_libreoffice = pytest.mark.skipif(
    find_soffice() is None,
    reason="LibreOffice is not installed; DOCX layout cannot be produced",
)


def _resume_with_table_and_header(path) -> None:
    document = docx.Document()
    document.sections[0].header.paragraphs[0].text = "Jane Doe | jane@example.com | +49 151 2345678"
    document.add_paragraph("Jane Doe")
    document.add_paragraph("Experience")
    document.add_paragraph("Senior Engineer at Acme Corp, 2021-Present")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Skills"
    table.rows[0].cells[1].text = "Python, SQL, Docker"
    document.save(str(path))


def test_convert_docx_returns_none_when_libreoffice_is_missing(tmp_path, monkeypatch):
    """Without a layout engine the caller must get None, not an exception:
    the app falls back to the text-only view rather than failing the upload.
    """
    import ats_xray.docx_render as module

    monkeypatch.setattr(module, "find_soffice", lambda: None)
    path = tmp_path / "resume.docx"
    docx.Document().save(str(path))

    assert convert_docx_to_pdf(str(path), str(tmp_path)) is None


def test_analyze_docx_without_libreoffice_still_reports_findings(tmp_path, monkeypatch):
    import ats_xray.docx_render as module

    monkeypatch.setattr(module, "find_soffice", lambda: None)
    path = tmp_path / "resume.docx"
    _resume_with_table_and_header(path)

    result = analyze_path(str(path), render=True)

    assert result.rendered_pages == []
    assert any(f.rule.id == "docx_table_content" for f in result.findings)


@needs_libreoffice
def test_convert_docx_produces_a_pdf(tmp_path):
    path = tmp_path / "resume.docx"
    _resume_with_table_and_header(path)

    converted = convert_docx_to_pdf(str(path), str(tmp_path))

    assert converted is not None
    assert converted.endswith(".pdf")
    naive, _ = extract_text(converted)
    assert "Jane Doe" in naive


@needs_libreoffice
def test_docx_findings_get_regions_on_the_rendered_page(tmp_path):
    path = tmp_path / "resume.docx"
    _resume_with_table_and_header(path)

    result = analyze_path(str(path), render=True)

    assert result.rendered_pages, "expected a rendered page for a DOCX"
    table_finding = next(f for f in result.findings if f.rule.id == "docx_table_content")
    header_finding = next(f for f in result.findings if f.rule.id == "docx_header_footer_content")
    assert table_finding.regions
    assert header_finding.regions


@needs_libreoffice
def test_header_text_is_not_matched_against_the_body(tmp_path):
    """Regression: the body line "Jane Doe" is a substring of the header
    "Jane Doe | jane@example.com | ...", and was being boxed as header
    content. Only the header itself should be marked.
    """
    path = tmp_path / "resume.docx"
    _resume_with_table_and_header(path)

    result = analyze_path(str(path), render=True)
    header_finding = next(f for f in result.findings if f.rule.id == "docx_header_footer_content")

    assert len(header_finding.regions) == 1, (
        f"expected only the header itself to be boxed, got {len(header_finding.regions)} regions"
    )


def test_find_text_regions_matches_a_snippet_inside_a_longer_line(tmp_path):
    """A table cell reading "Skills" really is present in the rendered row
    "Skills   Python, SQL, Docker", so a short snippet inside a line counts.
    """
    pdf_path = tmp_path / "page.pdf"
    c = canvas.Canvas(str(pdf_path), pagesize=(400, 200))
    c.setFont("Helvetica", 11)
    c.drawString(30, 150, "Skills   Python, SQL, Docker")
    c.drawString(30, 120, "Unrelated line of text here")
    c.save()

    regions = find_text_regions(str(pdf_path), ["Skills"])

    assert len(regions) == 1
    assert regions[0].top < 60


def test_find_text_regions_ignores_a_short_line_inside_a_longer_snippet(tmp_path):
    """The inverse direction is how rewrapped text is found, but a two-word
    line matching a long snippet is coincidence, not a wrap.
    """
    pdf_path = tmp_path / "page.pdf"
    c = canvas.Canvas(str(pdf_path), pagesize=(400, 200))
    c.setFont("Helvetica", 11)
    c.drawString(30, 150, "Jane Doe")
    c.save()

    regions = find_text_regions(str(pdf_path), ["Jane Doe | jane@example.com | +49 151 2345678"])

    assert regions == []


def test_attach_docx_regions_leaves_findings_alone_when_nothing_matches(tmp_path):
    """A finding whose content the layout engine did not place on the page
    keeps its evidence and simply has no box.
    """
    docx_path = tmp_path / "resume.docx"
    _resume_with_table_and_header(docx_path)

    empty_pdf = tmp_path / "empty.pdf"
    c = canvas.Canvas(str(empty_pdf), pagesize=(300, 300))
    c.setFont("Helvetica", 10)
    c.drawString(20, 200, "nothing in common")
    c.save()

    naive, aware = extract_text(str(docx_path))
    findings = run_rules(str(docx_path), naive, aware)

    located = attach_docx_regions(
        str(empty_pdf),
        findings,
        analyze_structure(str(docx_path)),
        build_field_report(aware),
        build_field_report(naive),
    )

    assert len(located) == len(findings)
    assert all(f.regions == () for f in located)
