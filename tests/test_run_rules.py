import docx
from reportlab.pdfgen import canvas

from ats_xray.docx_extract import extract_docx_full, extract_docx_naive
from ats_xray.engine import run_rules
from ats_xray.extract import extract_layout_aware, extract_naive


def test_run_rules_pdf_end_to_end_flags_column_hidden_section(tmp_path):
    """The same effect demonstrated elsewhere in the test suite: a section
    header split across two PDF columns is readable layout-aware but
    disappears under naive extraction, so the rule engine should flag it.
    """
    pdf_path = tmp_path / "resume.pdf"
    c = canvas.Canvas(str(pdf_path), pagesize=(500, 300))
    c.setFont("Helvetica", 12)
    c.drawString(30, 270, "Jane Doe")
    c.drawString(30, 250, "Experience")
    c.drawString(280, 270, "jane@example.com")
    c.drawString(280, 250, "+1 555 123 4567")
    c.save()

    naive_text = extract_naive(str(pdf_path))
    aware_text = extract_layout_aware(str(pdf_path))
    findings = run_rules(str(pdf_path), naive_text, aware_text)

    ids = {f.rule.id for f in findings}
    assert "section_missing_under_naive_parsing" in ids


def test_run_rules_docx_end_to_end_flags_table_content(tmp_path):
    path = tmp_path / "resume.docx"
    document = docx.Document()
    document.add_paragraph("Jane Doe")
    document.add_paragraph("jane@example.com, +1 555 123 4567")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Skills"
    table.rows[0].cells[1].text = "Python, SQL"
    document.save(str(path))

    naive_text = extract_docx_naive(str(path))
    aware_text = extract_docx_full(str(path))
    findings = run_rules(str(path), naive_text, aware_text)

    ids = {f.rule.id for f in findings}
    assert "docx_table_content" in ids


def test_run_rules_clean_resume_triggers_nothing(tmp_path):
    pdf_path = tmp_path / "resume.pdf"
    c = canvas.Canvas(str(pdf_path), pagesize=(300, 400))
    c.setFont("Helvetica", 12)
    c.drawString(30, 370, "Jane Doe")
    c.drawString(30, 350, "jane@example.com +1 555 123 4567")
    c.drawString(30, 300, "Experience")
    c.drawString(30, 280, "Senior Engineer at Acme")
    c.save()

    naive_text = extract_naive(str(pdf_path))
    aware_text = extract_layout_aware(str(pdf_path))
    findings = run_rules(str(pdf_path), naive_text, aware_text)

    assert findings == []
