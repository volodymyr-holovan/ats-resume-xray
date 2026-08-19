"""Finds DOCX content that lives outside the normal paragraph flow:
header/footer parts and text boxes.

Why this matters: both are common building blocks of visually appealing
resume templates, and both are common, well-documented causes of content
disappearing entirely during ATS parsing.

- **Headers/footers** live in separate XML parts (``word/header1.xml``,
  ``word/footer1.xml``) referenced from section properties, not inside
  ``document.xml``'s body. A reader that only walks the document body —
  including our own "full" extractor in ``docx_extract.py`` — never sees them.
- **Text boxes** nest their paragraphs inside a ``w:txbxContent`` element,
  itself buried inside a drawing anchor within a run. They don't appear as
  ordinary sibling paragraphs at the body level either, so a body-order walk
  misses them too. Sidebar sections built with text boxes are one of the
  most frequently cited reasons a resume "disappears" in an ATS.
"""

from typing import Any

import docx
from docx.oxml.ns import qn


def extract_docx_headers_footers(docx_path: str) -> dict:
    """Return ``{"headers": [str, ...], "footers": [str, ...]}`` with the
    non-empty text of each section's header/footer. Sections whose
    header/footer is linked to the previous section's (i.e. not distinct
    content) are skipped to avoid duplicate entries.
    """
    document = docx.Document(docx_path)
    headers = []
    footers = []

    for section in document.sections:
        if not section.header.is_linked_to_previous:
            text = _paragraphs_text(section.header)
            if text:
                headers.append(text)
        if not section.footer.is_linked_to_previous:
            text = _paragraphs_text(section.footer)
            if text:
                footers.append(text)

    return {"headers": headers, "footers": footers}


def has_table_content(docx_path: str) -> bool:
    """Return True if any DOCX table cell holds non-empty text.

    Tables are a common way to build a visually clean, side-by-side resume
    layout — and a common, well-documented way for that layout to break:
    many parsers flatten table rows in a way that scrambles which value
    belongs to which label, or skip table content entirely.
    """
    document = docx.Document(docx_path)
    return any(cell.text.strip() for table in document.tables for row in table.rows for cell in row.cells)


def find_docx_text_box_content(docx_path: str) -> list[str]:
    """Return the text found inside each Word text box (``w:txbxContent``)
    in the document body, in document order. Empty text boxes are skipped.
    """
    document = docx.Document(docx_path)
    contents = []

    for text_box in document.element.body.iter(qn("w:txbxContent")):
        text = "\n".join(
            "".join(node.text or "" for node in paragraph.iter(qn("w:t")))
            for paragraph in text_box.iter(qn("w:p"))
        ).strip()
        if text:
            contents.append(text)

    return contents


def _paragraphs_text(header_or_footer: Any) -> str:
    """``header_or_footer`` is a python-docx ``_Header``/``_Footer`` object;
    typed as ``Any`` since python-docx doesn't export a public type for it.
    """
    return "\n".join(p.text for p in header_or_footer.paragraphs if p.text.strip())
