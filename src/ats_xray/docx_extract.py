"""DOCX text extraction, in two flavors.

``extract_docx_naive`` only reads ``document.paragraphs`` — a common
shortcut in simple parsers. Any text placed inside a Word **table** (a
popular way to build a two-column resume) never appears in ``.paragraphs``
at all, so it is silently dropped.

``extract_docx_full`` walks the document body in true XML order, handling
paragraphs and tables as they actually appear in the file, so nothing
inside a table is lost.
"""

import docx
from docx.oxml.ns import qn
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph


def extract_docx_naive(docx_path: str) -> str:
    document = docx.Document(docx_path)
    return "\n".join(p.text for p in document.paragraphs if p.text.strip())


def extract_docx_full(docx_path: str) -> str:
    document = docx.Document(docx_path)
    return "\n".join(_blocks(document.element.body, document))


def _blocks(element, parent) -> list[str]:
    """Text of one container's children, in the order Word stores them."""
    parts: list[str] = []
    for child in element.iterchildren():
        if child.tag == qn("w:p"):
            text = Paragraph(child, parent).text.strip()
            if text:
                parts.append(text)
        elif child.tag == qn("w:tbl"):
            parts.extend(_table_rows(Table(child, parent), parent))
    return parts


def _table_rows(table: Table, parent) -> list[str]:
    """One line per row, cells joined with a pipe.

    ``row.cells`` reports a merged cell once for every grid column it
    covers, so a banner spanning two columns comes back twice on its own
    row and a merged sidebar comes back again in every row it reaches
    down through. A two-column CV built the way designer templates build
    one produced its contact block four times over.

    Merged cells share one underlying ``w:tc`` element, so remembering the
    elements already emitted -- for the whole table, not just the row --
    collapses both directions. Identity rather than text, because two
    genuinely separate cells may legitimately say the same thing.

    The elements are kept, not their ``id()``. lxml builds a proxy object
    on each access and frees it when the last reference goes, and CPython
    hands the freed address to the next proxy: with ``id()`` a merged
    sidebar still repeated and an unrelated cell was dropped as a
    duplicate. Holding the element keeps its proxy alive, which is what
    makes identity mean anything.
    """
    seen = set()
    lines: list[str] = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            if cell._tc in seen:
                continue
            seen.add(cell._tc)
            text = _cell_text(cell)
            if text:
                cells.append(text)
        if cells:
            lines.append(" | ".join(cells))
    return lines


def _cell_text(cell: _Cell) -> str:
    """A cell's own paragraphs and any table nested inside it.

    ``cell.text`` would cover the paragraphs and quietly drop a nested
    table, which is how the inner grid of a template built as a table
    inside a table disappears.
    """
    return "\n".join(_blocks(cell._tc, cell))
